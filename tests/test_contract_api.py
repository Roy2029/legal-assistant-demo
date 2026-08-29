"""合同审查 API 测试：上传→列表→审查→报告→下载还原→删除。"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from fastapi.testclient import TestClient
from server.main import app

client = TestClient(app)


def _make_contract(path: Path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("甲方：张三，电话13800138000。")
    doc.add_paragraph("合同价款为50万元，逾期按日千分之一支付违约金。")
    doc.add_paragraph("双方约定逾期审核视为认可送审价。")
    doc.add_paragraph("争议由守约方所在地法院管辖。")
    doc.add_paragraph("依据民法典第585条处理。")
    doc.save(str(path))


def test_contract_flow(tmp_path):
    src = tmp_path / "施工合同.docx"
    _make_contract(src)

    # 上传
    with open(src, "rb") as f:
        r = client.post("/api/contracts/upload", files={"files": ("施工合同.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200 and r.json()["ok"], r.text
    uploaded = r.json()["data"][0]
    assert uploaded["ok"] is True
    cid = uploaded["contract_id"]

    # 列表
    r = client.get("/api/contracts")
    assert r.json()["ok"]
    assert any(c["contract_id"] == cid for c in r.json()["data"])

    # 在线编辑
    r = client.get(f"/api/contracts/{cid}/files")
    assert r.json()["ok"] and len(r.json()["data"]) >= 1
    files = r.json()["data"]
    r = client.get(f"/api/contracts/{cid}/content", params={"file": files[0]})
    assert r.status_code == 200 and r.json()["ok"], r.text
    r = client.put(f"/api/contracts/{cid}/content", json={"file": files[0], "content": "编辑后的合同内容。逾期按日千分之一支付违约金。"})
    assert r.status_code == 200 and r.json()["ok"], r.text
    r = client.get(f"/api/contracts/{cid}/files")
    assert any(f.endswith("_编辑版.md") for f in r.json()["data"])

    # 审查
    r = client.post(f"/api/contracts/{cid}/review")
    assert r.status_code == 200 and r.json()["ok"], r.text
    d = r.json()
    assert d["risk_count"] >= 2, f"应命中至少2条内置规则，实际 {d['risk_count']}"

    # 报告
    r = client.get(f"/api/contracts/{cid}/report")
    assert r.status_code == 200

    # 下载还原版
    r = client.get(f"/api/contracts/{cid}/download", params={"kind": "restored"})
    assert r.status_code == 200
    text = r.content.decode("utf-8")
    assert "张三" in text, "还原版应包含真实姓名"

    # 下载批注版
    r = client.get(f"/api/contracts/{cid}/download", params={"kind": "annotated"})
    assert r.status_code == 200, r.text
    import docx as dx
    from tempfile import NamedTemporaryFile
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(r.content)
        tmp_path = Path(tmp.name)
    d = dx.Document(str(tmp_path))
    all_text = "\n".join(p.text for p in d.paragraphs)
    assert "风险提示" in all_text, "批注版应包含风险提示"
    tmp_path.unlink(missing_ok=True)

    # 删除
    r = client.delete(f"/api/contracts/{cid}")
    assert r.json()["ok"]
    r = client.get("/api/contracts")
    assert not any(c["contract_id"] == cid for c in r.json()["data"])
    print("PASS contract_flow")


def test_upload_skill_md():
    skill_content = "# 用户 skill\n\n1. 先看合同\n2. 再查规则\n".encode("utf-8")
    r = client.post("/api/contracts/skills", files={"file": ("我的审查流程.md", skill_content, "text/markdown")})
    assert r.status_code == 200 and r.json()["ok"], r.text
    assert r.json()["data"]["type"] == "skill"
    r = client.get("/api/contracts/skills")
    assert r.json()["ok"]
    assert "我的审查流程.md" in r.json()["data"]["skills"]
    # 清理
    Path("skills/contract_review/user_skills/我的审查流程.md").unlink(missing_ok=True)
    print("PASS upload_skill_md")


if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as d:
        test_contract_flow(Path(d))
    print("ALL CONTRACT API TESTS PASSED")


def test_configurable_mask_flow(tmp_path):
    import docx as dx
    src = tmp_path / "可配置脱敏测试.docx"
    doc = dx.Document()
    doc.add_paragraph("甲方：张三，电话13800138000，邮箱 zhangsan@example.com，身份证号110101199001011234。")
    doc.add_paragraph("乙方：华为技术有限公司，统一社会信用代码91110108MA01ABCD1X。")
    doc.add_paragraph("合同价款为50万元，逾期按日千分之一支付违约金。")
    doc.save(str(src))

    with open(src, "rb") as f:
        r = client.post("/api/contracts/upload", files={"files": ("可配置脱敏测试.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200 and r.json()["ok"], r.text
    cid = r.json()["data"][0]["contract_id"]

    # 版本列表应包含原版与初始脱敏版
    r = client.get(f"/api/contracts/{cid}/versions")
    assert r.status_code == 200 and r.json()["ok"]
    kinds = [v["kind"] for v in r.json()["data"]]
    assert "original" in kinds and "redacted" in kinds

    # 扫描应命中六类敏感信息
    r = client.get(f"/api/contracts/{cid}/scan")
    assert r.status_code == 200 and r.json()["ok"]
    items = r.json()["data"]["items"]
    cats = {i["category"] for i in items}
    assert {"person_name", "company_name", "credit_code", "phone", "email", "id_card"} <= cats

    # 按扫描清单确认脱敏（占位符）
    r = client.post(f"/api/contracts/{cid}/mask", json={"categories": ["person_name", "company_name", "credit_code", "phone", "email", "id_card"], "method": "placeholder", "items": items})
    assert r.status_code == 200 and r.json()["ok"], r.text
    md = r.json()["data"]
    assert md["masked_count"] == len(items)
    assert "张三" not in md["content"] and "华为技术有限公司" not in md["content"]

    # 映射配置可还原
    r = client.get(f"/api/contracts/{cid}/mask/mapping")
    assert r.json()["ok"]
    entries = r.json()["data"]["entries"]
    assert len(entries) == len(items)
    r = client.post(f"/api/contracts/{cid}/mask/restore", json={"file": md["file"], "entries": entries})
    assert r.status_code == 200 and r.json()["ok"]
    assert "张三" in r.json()["data"]["content"]
    assert "13800138000" in r.json()["data"]["content"]

    # 拖选手动片段加入清单
    r = client.post(f"/api/contracts/{cid}/mask/manual", json={"text": "特殊敏感句", "category": "manual"})
    assert r.status_code == 200 and r.json()["ok"]
    r = client.get(f"/api/contracts/{cid}/mask/manual")
    assert len(r.json()["data"]) >= 1

    # 重命名
    r = client.put(f"/api/contracts/{cid}", json={"original_name": "重命名合同.docx"})
    assert r.status_code == 200 and r.json()["ok"]

    # 删除
    r = client.delete(f"/api/contracts/{cid}")
    assert r.json()["ok"]
    print("PASS configurable_mask_flow")
