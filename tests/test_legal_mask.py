"""LegalMask vendored 脱敏/还原测试。"""
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from online_core.legal_mask import process_document, restore_text


def test_process_docx_and_restore(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("甲方：张三，身份证号110101199001011234，电话13800138000。")
    doc.add_paragraph("合同价款为50万元，逾期按日千分之一支付违约金。")
    doc.add_paragraph("依据民法典第580条处理。")
    src = tmp_path / "test_contract.docx"
    doc.save(str(src))

    out = tmp_path / "out"
    result = process_document(str(src), str(out))
    assert result.success, result.error
    assert result.output_file and Path(result.output_file).exists()
    assert result.mapping_file and Path(result.mapping_file).exists()

    mapping = json.loads(Path(result.mapping_file).read_text(encoding="utf-8"))
    assert mapping, "mapping 不应为空"

    cleaned = Path(result.output_file).read_text(encoding="utf-8") if result.output_file.endswith(".txt") else _docx_text(Path(result.output_file))
    assert "张三" not in cleaned, "脱敏后不应出现真实姓名"
    assert "民法典" in cleaned or "第580条" in cleaned or "__LAW_" in cleaned, "法条引用应被保护"

    restored = restore_text(cleaned, mapping)
    assert "张三" in restored or "13800138000" in restored, "还原后应包含真实信息"
    print("PASS legal_mask_process_and_restore")


def _docx_text(path):
    import docx
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as d:
        test_process_docx_and_restore(Path(d))
    print("ALL LEGAL MASK TESTS PASSED")
