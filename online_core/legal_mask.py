"""LegalMask core processor — vendored from https://github.com/RayBagel/LegalMask (MIT License).

Original file: core/processor.py
We vendor it locally for the contract review agent's reversible redaction.
Copyright (c) LegalMask authors. See data/tmp_selection/LegalMask/LICENSE for original license.
"""
"""
LegalMask core processor — local redaction engine, no external API.
Extracted and restructured from clean_local_v2.py.
"""

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Tuple, Dict


# ---------------------------------------------------------------------------
# Public data types
# ---------------------------------------------------------------------------

@dataclass
class ProcessResult:
    success: bool
    input_file: str
    output_file: str | None = None
    mapping_file: str | None = None
    mapping_dict: dict = field(default_factory=dict)
    error: str | None = None
    entity_count: int = 0


# ---------------------------------------------------------------------------
# Internal: core redaction logic (ported from clean_sensitive_info)
# ---------------------------------------------------------------------------

def _clean_sensitive_info(text: str) -> Tuple[str, Dict[str, str]]:
    """脱敏处理，返回 (脱敏后文本, 占位符→原值 mapping_dict)。"""

    mapping_dict: Dict[str, str] = {}
    case_counter = 1

    # ── Layer 0: protect statute references ──────────────────────────────
    law_refs = []
    law_counter = 1

    law_pattern = (
        r"第"
        r"(?:[零一二三四五六七八九十百千万]+|\d+)"
        r"(?:条|款|项|)?"
        r"(?:之[零一二三四五六七八九十百千万\d]+)?"
        r"(?:第(?:[零一二三四五六七八九十百千万]+|\d+)(?:条|款|项))*"
    )

    def replace_law(match: re.Match) -> str:
        nonlocal law_counter
        ref = match.group(0)
        marker = f"__LAW_{law_counter}__"
        law_refs.append((marker, ref))
        law_counter += 1
        return marker

    text = re.sub(law_pattern, replace_law, text)

    # ── Layer 0.5: collect party surnames → assign codes ─────────────────
    temp_text_for_scan = text
    surnames: set = set()
    full_names: Dict[str, str] = {}

    legal_roles = (
        r"(?:"
        r"原告|被告|第三人|上诉人|被上诉人|申请人|被申请人\b"
        r"|再审申请人|再审被申请人|申请执行人|被执行人\b"
        r"|法定代表人|负责人|执行事务合伙人\b"
        r"|委托诉讼代理人|委托代理人|诉讼代理人|辩护人\b"
        r"|联系人|经办人|授权代表|经纪人|代理人\b"
        r"|反诉原告|反诉被告|本诉原告|本诉被告\b"
        r"|债权人|债务人|担保人|保证人|抵押人|出质人\b"
        r"|买方|卖方|出租方|承租方|发包方|承包方\b"
        r"|甲方|乙方|丙方|丁方\b"
        r")"
    )

    pattern1 = rf"({legal_roles})([\s：:是为]*)([一-龥])([一-龥]{{0,3}})\b"
    for m in re.finditer(pattern1, temp_text_for_scan):
        if m.group(3):
            surnames.add(m.group(3))
            full_name = m.group(3) + m.group(4)
            if m.group(3) not in full_names:
                full_names[m.group(3)] = full_name

    pattern2 = r"([一-龥])(?:某{1,2})\b"
    for m in re.finditer(pattern2, temp_text_for_scan):
        if m.group(1):
            surnames.add(m.group(1))
            if m.group(1) not in full_names:
                full_names[m.group(1)] = m.group(0)

    pattern3 = r"(?:姓名|名字|称)[\s：:]*([一-龥])(?:[一-龥]{0,3})\b"
    for m in re.finditer(pattern3, temp_text_for_scan):
        if m.group(1):
            surnames.add(m.group(1))
            fn_match = re.search(
                r"(?:姓名|名字|称)[\s：:]*([一-龥]{1,4})\b", m.group(0)
            )
            if fn_match and m.group(1) not in full_names:
                full_names[m.group(1)] = fn_match.group(1)

    party_codes = [
        "甲方", "乙方", "丙方", "丁方", "戊方",
        "己方", "庚方", "辛方", "壬方", "癸方",
    ]
    surname_to_code: Dict[str, str] = {}
    for i, surname in enumerate(sorted(surnames)):
        code = party_codes[i] if i < len(party_codes) else f"方{i + 1}"
        surname_to_code[surname] = code
        if surname in full_names:
            mapping_dict[f"{{{{PARTY_{code}}}}}"] = full_names[surname]

    # ── Layer 1: unique identifiers（带编号，可精确还原）─────────────────
    def _replace_numbered(text, pattern, base):
        counter = 1
        out = []
        pos = 0
        for m in re.finditer(pattern, text):
            original = m.group(0)
            placeholder = f"{{{{{base}_{counter}}}}}"
            mapping_dict[placeholder] = original
            out.append(text[pos:m.start()])
            out.append(placeholder)
            pos = m.end()
            counter += 1
        out.append(text[pos:])
        return "".join(out)

    text = _replace_numbered(text, r"(?<!\d)\d{17}[\dXx](?!\d)", "ID")
    text = _replace_numbered(text, r"(?<!\d)1[3-9]\d{9}(?!\d)", "PHONE")
    text = _replace_numbered(text, r"(?<![0-9A-Z])[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}(?![0-9A-Z])", "CREDIT_CODE")

    # ── Layer 2: amounts（带编号，可精确还原）───────────────────────────
    text = _replace_numbered(text, r"[\d,]+\.?\d*\s*(?:万元|亿元|元|美元|英镑|港币)", "AMOUNT")
    text = _replace_numbered(text, r"\d+\.?\d*\s*(?:万元|亿元|元)", "AMOUNT")
    text = _replace_numbered(text, r"\d+\.?\d*\s*元/(?:月|日|年)", "AMOUNT_PERIOD")
    text = _replace_numbered(text, r"\b\d+\.?\d*%\b", "PERCENT")
    text = re.sub(r"年利率\s*\d+\.?\d*%", "年利率{{PERCENT}}", text)
    text = re.sub(r"日利率\s*\d+\.?\d*%", "日利率{{PERCENT}}", text)

    # ── Layer 3: case numbers & reference IDs ─────────────────────────────
    def replace_case_number(match: re.Match) -> str:
        nonlocal case_counter
        original = match.group(0)
        placeholder = f"{{{{CASE_NUMBER_{case_counter}}}}}"
        mapping_dict[placeholder] = original
        case_counter += 1
        return placeholder

    text = re.sub(
        r"[（(]\d{4}[）)]"
        r"[一-龥]{0,4}"
        r"\d+"
        r"[一-龥]*"
        r"\d+"
        r"[a-zA-Z一-龥]*"
        r"号",
        replace_case_number,
        text,
    )

    text = re.sub(r"入库编号[\s：:]*\d{4}-\d+-\d+-\d+-\d+", "入库编号：{{REFERENCE_NUMBER}}", text)
    text = re.sub(r"档案编号[\s：:]*[\w-]+", "档案编号：{{REFERENCE_NUMBER}}", text)
    text = re.sub(r"案件编号[\s：:]*[\w-]+", "案件编号：{{REFERENCE_NUMBER}}", text)

    # ── Layer 4: institution names ────────────────────────────────────────
    text = re.sub(r"[一-龥]+?(?:高级|中级|基层|专门)?(?:人民)?法院\b", "{{COURT}}", text)
    text = re.sub(
        r"[一-龥]+?(?:银行|保险|证券|信托|基金|期货)(?:公司|集团)?\b",
        "{{FINANCIAL_INST}}",
        text,
    )
    text = re.sub(
        r"[一-龥]+?(?:大学|学院|学校|中学|小学|幼儿园|培训)(?:机构)?\b",
        "{{SCHOOL}}",
        text,
    )
    text = re.sub(
        r"[一-龥]+?(?:政府|委员会|局|厅|处|科|办|部)(?:机关)?\b",
        "{{GOVERNMENT}}",
        text,
    )
    text = re.sub(r"某[一-龥]{1,10}(?:公司|企业|集团|厂|院|校)", "{{ORGANIZATION}}", text)
    org_suffixes = r"(?:公司|企业|集团|厂|院|所|校|中心|站|社|店|馆|部|处|科|室)"
    text = re.sub(r"[一-龥]{2,20}?" + org_suffixes, "{{ORGANIZATION}}", text)

    text = re.sub(r"账户名称[\s：:]*[一-龥]{2,20}", "账户名称：{{ACCOUNT_NAME}}", text)
    text = re.sub(r"账号[\s：:]*\d{10,20}", "账号：{{ACCOUNT_NUMBER}}", text)
    text = re.sub(
        r"开户行[\s：:]*[一-龥]+(?:银行|信用社|储蓄所)",
        "开户行：{{BANK_BRANCH}}",
        text,
    )
    text = re.sub(r"合同编号[\s：:]*[\w-]+\d+", "合同编号：{{CONTRACT_NUMBER}}", text)
    text = re.sub(r"签订日期[\s：:]*\d{4}年\d{1,2}月\d{1,2}日", "签订日期：{{DATE}}", text)
    text = re.sub(
        r"(审判长|审判员|书记员)[\s：:]*[一-龥]{2,4}", r"\1：{{JUDGE}}", text
    )
    text = re.sub(r"审理日期[\s：:]*\d{4}年\d{1,2}月\d{1,2}日", "审理日期：{{DATE}}", text)
    text = re.sub(r"判决日期[\s：:]*\d{4}年\d{1,2}月\d{1,2}日", "判决日期：{{DATE}}", text)
    text = re.sub(r"(?<!\d)\d{16,19}(?!\d)", "{{BANK}}", text)

    # ── Layer 5: person names (using surname codes) ───────────────────────
    def replace_by_surname(surname: str) -> str:
        if surname in surname_to_code:
            return f"{{{{PARTY_{surname_to_code[surname]}}}}}"
        return "{{PARTY}}"

    def replace_pattern1(match: re.Match) -> str:
        role = match.group(1)
        separator = match.group(2) if match.group(2) else ""
        surname = match.group(3)
        return f"{role}{separator}{replace_by_surname(surname)}"

    text = re.sub(
        rf"({legal_roles})([\s：:是为]*)([一-龥])([一-龥]{{0,3}})\b",
        replace_pattern1,
        text,
    )

    def replace_pattern2(match: re.Match) -> str:
        label = match.group(1)
        surname = match.group(2)
        return f"{label}：{replace_by_surname(surname)}"

    text = re.sub(
        r"(姓名|名字|称)[\s：:]*([一-龥])(?:[一-龥]{0,3})\b",
        replace_pattern2,
        text,
    )

    text = re.sub(
        r"([一-龥])某某",
        lambda m: replace_by_surname(m.group(1)),
        text,
    )
    text = re.sub(
        r"([一-龥])某",
        lambda m: replace_by_surname(m.group(1)),
        text,
    )

    context_patterns = r"(?:与|向|对|由|为|和|跟|同|被|让|给)"

    def replace_context(match: re.Match) -> str:
        prefix = match.group(1)
        name = match.group(2)
        surname = name[0] if name else ""
        return f"{prefix}{replace_by_surname(surname)}"

    text = re.sub(
        rf"({context_patterns})([一-龥]{{2,4}})(?:签订|起诉|上诉|辩称|诉称|称|表示|认为|主张)",
        replace_context,
        text,
    )

    # ── Layer 6: dates and addresses ──────────────────────────────────────
    text = re.sub(r"\d{4}年\d{1,2}月\d{1,2}日", "{{DATE}}", text)
    text = re.sub(r"\d{4}年\d{1,2}月", "{{DATE}}", text)
    text = re.sub(r"\d{4}-\d{1,2}-\d{1,2}", "{{DATE}}", text)
    text = re.sub(r"\d{4}/\d{1,2}/\d{1,2}", "{{DATE}}", text)

    address_keywords = r"(?:省|市|区|县|镇|乡|村|街道|路|街|巷|弄|号|栋|幢|楼|座|层|室|房|院)"

    protected_patterns = [
        (r"(?<!\w)身份证号[\s：:]*", "__PROTECTED_ID__"),
        (r"(?<!\w)户口[\s：:]*", "__PROTECTED_HOUSEHOLD__"),
        (r"(?<!\w)银行[\s：:]*", "__PROTECTED_BANK__"),
        (r"(?<!\w)信用卡[\s：:]*", "__PROTECTED_CREDIT__"),
        (r"(?<!\w)合同[\s：:]*", "__PROTECTED_CONTRACT__"),
        (r"(?<!\w)案件[\s：:]*", "__PROTECTED_CASE__"),
        (r"(?<!\w)档案[\s：:]*", "__PROTECTED_ARCHIVE__"),
        (r"(?<!\w)入库[\s：:]*", "__PROTECTED_STORAGE__"),
    ]

    temp_text = text
    replacements_made = []
    for pattern, temp_marker in protected_patterns:
        matches = list(re.finditer(pattern, temp_text))
        if matches:
            for match in matches:
                replacements_made.append((match.start(), match.end(), match.group()))
            temp_text = re.sub(pattern, temp_marker, temp_text)

    temp_text = re.sub(
        rf"[一-龥]{{2,30}}?(?:{address_keywords})[一-龥\d]{{0,50}}?(?:{address_keywords})?\b",
        "{{ADDRESS}}",
        temp_text,
    )

    for _start, _end, original in reversed(replacements_made):
        marker_map = {
            "__PROTECTED_ID__": original,
            "__PROTECTED_HOUSEHOLD__": original,
            "__PROTECTED_BANK__": original,
            "__PROTECTED_CREDIT__": original,
            "__PROTECTED_CONTRACT__": original,
            "__PROTECTED_CASE__": original,
            "__PROTECTED_ARCHIVE__": original,
            "__PROTECTED_STORAGE__": original,
        }
        for marker, orig in marker_map.items():
            if marker in temp_text:
                temp_text = temp_text.replace(marker, orig, 1)

    text = temp_text
    text = re.sub(r"\b\d{6}\b(?=.*?(?:邮编|邮政编码|Postal))", "{{POSTAL_CODE}}", text)

    # ── Layer 7: miscellaneous ────────────────────────────────────────────
    text = re.sub(r"[一-龥][A-Z][A-Z0-9]{5,6}\b", "{{LICENSE_PLATE}}", text)
    text = re.sub(r"[\w.-]+@[\w.-]+\.\w+", "{{EMAIL}}", text)
    text = re.sub(r"https?://[\w.-]+(?:/[\w./%-]*)?", "{{URL}}", text)
    text = re.sub(r"(?:QQ|微信|微信号)[\s：:]*[\w-]{6,20}", "{{SOCIAL_ACCOUNT}}", text)

    # ── Layer 8: cleanup ──────────────────────────────────────────────────
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"  +", " ", text)
    text = re.sub(r"\t+", " ", text)
    text = re.sub(r"\{\{PARTY\}\}\s*\{\{PARTY\}\}", "{{PARTY}}", text)
    text = re.sub(r"\{\{AMOUNT\}\}\s*\{\{AMOUNT\}\}", "{{AMOUNT}}", text)
    text = re.sub(r"\{\{DATE\}\}\s*\{\{DATE\}\}", "{{DATE}}", text)
    text = re.sub(r"\{\{JUDGE\}\}\s*\{\{JUDGE\}\}", "{{JUDGE}}", text)

    # ── Layer 9: restore statute references ───────────────────────────────
    for marker, ref in law_refs:
        text = text.replace(marker, ref)

    return text, mapping_dict


def _count_entities(mapping_dict: Dict[str, str], redacted_text: str) -> int:
    """Count distinct placeholders that actually appear in the redacted text."""
    static_tags = re.findall(r"\{\{[A-Z_]+(?:_\d+)?\}\}", redacted_text)
    all_keys = set(mapping_dict.keys()) | set(static_tags)
    return len(all_keys)


# ---------------------------------------------------------------------------
# Internal: file extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_plain(filepath: Path) -> str:
    """Extract text from .txt or .md."""
    return filepath.read_text(encoding="utf-8", errors="ignore")


def _extract_text_docx(filepath: Path) -> str:
    """Extract plain text from a .docx file."""
    import docx  # python-docx
    doc = docx.Document(str(filepath))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _extract_text_pdf_fitz(filepath: Path) -> str:
    """Extract text from a digital PDF using PyMuPDF."""
    import fitz
    doc = fitz.open(str(filepath))
    parts = [doc[i].get_text() for i in range(len(doc)) if doc[i].get_text().strip()]
    doc.close()
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Internal: per-run application (uses pre-built global mapping)
# ---------------------------------------------------------------------------

def _apply_run_redactions(text: str, reverse_map: Dict[str, str]) -> str:
    """Apply global mapping + generic patterns to a single run's text."""
    # Apply known mappings (party names, case numbers) first
    for original, placeholder in reverse_map.items():
        text = text.replace(original, placeholder)

    # Protect law refs, apply generic patterns, restore
    law_refs: list = []
    law_counter = [1]
    law_pattern = (
        r"第(?:[零一二三四五六七八九十百千万]+|\d+)"
        r"(?:条|款|项|)?(?:之[零一二三四五六七八九十百千万\d]+)?"
        r"(?:第(?:[零一二三四五六七八九十百千万]+|\d+)(?:条|款|项))*"
    )
    def _protect_law(m: re.Match) -> str:
        marker = f"__LAW_{law_counter[0]}__"
        law_refs.append((marker, m.group(0)))
        law_counter[0] += 1
        return marker
    text = re.sub(law_pattern, _protect_law, text)

    text = re.sub(r"(?<!\d)\d{17}[\dXx](?!\d)", "{{ID}}", text)
    text = re.sub(r"(?<!\d)1[3-9]\d{9}(?!\d)", "{{PHONE}}", text)
    text = re.sub(
        r"(?<![0-9A-Z])[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}(?![0-9A-Z])",
        "{{CREDIT_CODE}}", text,
    )
    text = re.sub(r"[\d,]+\.?\d*\s*(?:万元|亿元|元|美元|英镑|港币)", "{{AMOUNT}}", text)
    text = re.sub(r"\d+\.?\d*\s*(?:万元|亿元|元)", "{{AMOUNT}}", text)

    for marker, ref in law_refs:
        text = text.replace(marker, ref)
    return text


# ---------------------------------------------------------------------------
# Internal: format-preserving redaction
# ---------------------------------------------------------------------------

def _process_docx(input_path: Path, output_path: Path) -> Dict[str, str]:
    """Run-by-run redaction on a DOCX, preserving formatting.
    Builds global mapping once from full text, then applies per-run for consistency.
    """
    import docx

    full_text = _extract_text_docx(input_path)
    _, mapping_dict = _clean_sensitive_info(full_text)
    # reverse: original_value → placeholder so we can do string replace per-run
    reverse_map = {v: k for k, v in mapping_dict.items()}

    doc = docx.Document(str(input_path))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = _apply_run_redactions(run.text, reverse_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = _apply_run_redactions(run.text, reverse_map)

    doc.save(str(output_path))
    return mapping_dict


def _process_pdf_fitz(input_path: Path, output_path: Path) -> Dict[str, str]:
    """Page-by-page redaction on a PDF using PyMuPDF."""
    import fitz

    full_text = _extract_text_pdf_fitz(input_path)
    _, mapping_dict = _clean_sensitive_info(full_text)

    doc = fitz.open(str(input_path))
    new_doc = fitz.open()

    for page_num in range(len(doc)):
        page = doc[page_num]
        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[6] == 0:
                x0, y0, x1, y1, block_text, block_num, block_type = block
                cleaned_block, _ = _clean_sensitive_info(block_text)
                new_page.insert_text((x0, y0), cleaned_block, fontsize=10, color=(0, 0, 0))

    new_doc.save(str(output_path))
    new_doc.close()
    doc.close()
    return mapping_dict


def _process_pdf_fallback(input_path: Path, output_path: Path) -> Dict[str, str]:
    """Fallback: extract text with pdfplumber, save redacted .txt."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(input_path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t.strip())
        raw = "\n\n".join(parts)
    except Exception:
        raw = ""

    cleaned, mapping_dict = _clean_sensitive_info(raw)
    txt_output = output_path.with_suffix(".txt")
    txt_output.write_text(cleaned, encoding="utf-8")
    return mapping_dict


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def process_document(input_path: str | Path, output_dir: str | Path) -> ProcessResult:
    """处理单个DOCX或数字PDF文件。
    - 脱敏后文件保存到 output_dir/<原文件名>_cleaned.<ext>
    - mapping保存到 output_dir/<原文件名>_mapping.json
    - 返回ProcessResult
    """
    input_path = Path(input_path)
    output_dir = Path(output_dir)

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        return ProcessResult(
            success=False,
            input_file=str(input_path),
            error=f"无法创建输出目录: {exc}",
        )

    suffix = input_path.suffix.lower()
    stem = input_path.stem
    output_file = output_dir / f"{stem}_cleaned{suffix}"
    mapping_file = output_dir / f"{stem}_mapping.json"

    try:
        if suffix == ".docx":
            mapping_dict = _process_docx(input_path, output_file)
        elif suffix == ".pdf":
            try:
                mapping_dict = _process_pdf_fitz(input_path, output_file)
            except ImportError:
                output_file = output_dir / f"{stem}_cleaned.txt"
                mapping_dict = _process_pdf_fallback(input_path, output_file)
        elif suffix in {".txt", ".md"}:
            raw = _extract_text_plain(input_path)
            cleaned, mapping_dict = _clean_sensitive_info(raw)
            output_file = output_dir / f"{stem}_cleaned{suffix}"
            output_file.write_text(cleaned, encoding="utf-8")
        else:
            return ProcessResult(
                success=False,
                input_file=str(input_path),
                error=f"不支持的文件格式: {suffix}",
            )

        mapping_file.write_text(
            json.dumps(mapping_dict, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        entity_count = len(mapping_dict)

        return ProcessResult(
            success=True,
            input_file=str(input_path),
            output_file=str(output_file),
            mapping_file=str(mapping_file),
            mapping_dict=mapping_dict,
            entity_count=entity_count,
        )

    except Exception as exc:
        return ProcessResult(
            success=False,
            input_file=str(input_path),
            error=str(exc),
        )


def batch_process(input_paths: list, output_dir: str | Path) -> list:
    """批量处理多个文件，返回 list[ProcessResult]"""
    output_dir = Path(output_dir)
    return [process_document(p, output_dir) for p in input_paths]


def restore_text(ai_text: str, mapping_dict: dict) -> str:
    """把AI分析文字里的占位符（如{{PARTY_甲方}}）还原为真实姓名。
    遍历mapping_dict，把每个key替换为value。"""
    result = ai_text
    for placeholder, original in mapping_dict.items():
        result = result.replace(placeholder, original)
    return result


def load_mapping(mapping_json_path: str | Path) -> dict:
    """从JSON文件加载mapping_dict"""
    mapping_json_path = Path(mapping_json_path)
    with open(mapping_json_path, encoding="utf-8") as f:
        return json.load(f)
