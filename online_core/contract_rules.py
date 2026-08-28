"""合同审查规则引擎：内置 + 用户规则加载、文本扫描、报告生成。"""
from __future__ import annotations

import json
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
BUILTIN_RULES = PROJECT_ROOT / "skills" / "contract_review" / "rules.jsonl"
USER_RULES_DIR = PROJECT_ROOT / "skills" / "contract_review" / "user_rules"


def load_rules() -> list[dict]:
    rules = []
    for path in [BUILTIN_RULES, *sorted(USER_RULES_DIR.glob("*.jsonl"))]:
        if not path.exists():
            continue
        try:
            for line in path.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line:
                    continue
                rule = json.loads(line)
                if rule.get("status", "enabled") != "disabled":
                    rules.append(rule)
        except Exception:
            continue
    return rules


def scan_text(text: str, rules: list[dict] | None = None, file_name: str = "") -> list[dict]:
    rules = rules if rules is not None else load_rules()
    risks = []
    for rule in rules:
        patterns = (rule.get("trigger") or {}).get("patterns") or []
        for pat in patterns:
            if not pat:
                continue
            idx = text.find(pat)
            if idx >= 0:
                snippet = text[max(0, idx - 40):idx + 120].replace("\n", " ")
                risks.append({
                    "rule_id": rule.get("rule_id"),
                    "dimension": rule.get("dimension"),
                    "risk_level": rule.get("risk_level"),
                    "risk_desc": rule.get("risk_desc"),
                    "suggestion": rule.get("suggestion_template"),
                    "basis": rule.get("basis") or [],
                    "source": rule.get("source", ""),
                    "file": file_name,
                    "snippet": snippet,
                })
                break

    # 去重：同一 rule_id + file 只保留一条
    seen = set()
    deduped = []
    for r in risks:
        key = (r["rule_id"], r["file"])
        if key not in seen:
            seen.add(key)
            deduped.append(r)
    return deduped


def annotate_docx(input_path: str | Path, output_path: str | Path, rules: list[dict] | None = None) -> int:
    """在脱敏 DOCX 上生成批注版：命中风险的段落黄色高亮，段后插入红色风险提示。

    Returns: 批注数量
    """
    import docx
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import RGBColor

    rules = rules if rules is not None else load_rules()
    doc = docx.Document(str(input_path))
    notes = 0
    for para in list(doc.paragraphs):
        text = para.text or ""
        if not text.strip():
            continue
        risks = scan_text(text, rules)
        if not risks:
            continue
        # 高亮原段落
        for run in para.runs:
            try:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except Exception:
                pass
        # 段后插入风险提示
        note = doc.add_paragraph()
        tip = "【风险提示】" + "；".join(
            f"[{r['risk_level']}] {r['risk_desc']} → {r['suggestion']}" for r in risks
        )
        run = note.add_run(tip)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.font.bold = True
        # 把 note 从文末移动到当前段之后
        try:
            para._p.addnext(note._p)
        except Exception:
            pass
        notes += 1
    doc.save(str(output_path))
    return notes


def render_report(file_names: list[str], risks: list[dict]) -> str:
    high = [r for r in risks if r["risk_level"] == "high"]
    medium = [r for r in risks if r["risk_level"] == "medium"]
    low = [r for r in risks if r["risk_level"] == "low"]
    report = "# 合同审查报告\n\n"
    report += f"- 合同：{', '.join(file_names)}\n"
    report += f"- 风险总数：{len(risks)}（高 {len(high)} / 中 {len(medium)} / 低 {len(low)}）\n\n"
    report += "## 风险清单\n\n"
    if not risks:
        report += "未命中内置规则。\n"
    for r in risks:
        report += f"### [{r['risk_level']}] {r['dimension']} — {r['rule_id']}\n"
        report += f"- 说明：{r['risk_desc']}\n"
        report += f"- 建议：{r['suggestion']}\n"
        report += f"- 原文片段：{r['snippet']}\n"
        if r["basis"]:
            report += f"- 依据：{r['basis']}\n"
        report += "\n"
    report += "---\n本报告由规则引擎生成，仅供参考，使用前须经执业律师核阅。\n"
    return report
