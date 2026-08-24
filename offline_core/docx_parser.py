"""DOCX 解析器。

提供两个层次的解析组件：
- DocxParser：通用 DOCX 解析，基于 python-docx 原生样式（Heading 1/2/3）
- LegalDocxParser：法律文档专用，基于正则模式 + 字号检测
"""

import os
import re
import uuid
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.oxml.ns import qn

from .modules import Parser
from .data_model import (
    StructuredDocument,
    HeadingBlock,
    ParagraphBlock,
    TableBlock,
    ImageBlock,
)


# ── 中文数字转换 ──────────────────────────────────────────────

_CN_NUM = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
    "百": 100, "千": 1000, "零": 0,
}


def chinese_to_arabic(chinese: str) -> int:
    """中文数字 → int。例: '一'→1, '十二'→12, '一百二十'→120。"""
    result = 0
    temp = 0
    for ch in chinese:
        num = _CN_NUM.get(ch, 0)
        if num >= 10:
            temp = max(temp, 1) * num
            result += temp
            temp = 0
        else:
            temp = num
    return result + temp


# ── 正则模式（法律文档） ──────────────────────────────────────

_RE_CHAPTER = re.compile(r"^第([一二三四五六七八九十百零千]+)章\s*(.*)")
_RE_SECTION = re.compile(r"^第([一二三四五六七八九十百零千]+)节\s*(.*)")
_RE_ARTICLE = re.compile(r"^第([一二三四五六七八九十百零千]+)条\s*(.*)")
_RE_APPENDIX = re.compile(r"^附件([一二三四五六七八九十百零千]+)\s*(.*)")
_RE_PREFACE = re.compile(r"^序\s*言")
_RE_TOC = re.compile(r"目\s*录")


# ═══════════════════════════════════════════════════════════════
#  通用 DocxParser
# ═══════════════════════════════════════════════════════════════

class DocxParser(Parser):
    """通用 DOCX 解析器，基于 python-docx 原生样式检测标题。"""

    def __init__(
        self,
        extract_tables: bool = True,
        extract_images: bool = False,
        image_output_dir: str = "./docx_images",
    ):
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.image_output_dir = image_output_dir

    # ── 标题检测 ──────────────────────────────────────────────

    def _get_heading_level(self, para) -> Optional[int]:
        """从段落样式中提取标题级别。"""
        style = para.style
        if style is None:
            return None

        name = style.name or ""

        # 英文：Heading 1 → 1
        if name.startswith("Heading"):
            try:
                return int(name.split()[-1])
            except (ValueError, IndexError):
                return 1

        # 中文：标题 1 → 1
        zh_match = re.match(r"^标题\s*(\d+)", name)
        if zh_match:
            return int(zh_match.group(1))

        # 其他标题类样式
        lower = name.lower()
        if "title" in lower or "head" in lower:
            return 1

        return None

    # ── 段落解析 ──────────────────────────────────────────────

    def _parse_paragraph(self, para, order: int):
        """解析单个段落 → HeadingBlock 或 ParagraphBlock。"""
        text = para.text.strip()
        if not text:
            return None

        heading_level = self._get_heading_level(para)
        if heading_level is not None:
            return HeadingBlock(order=order, content=text, level=heading_level)

        return ParagraphBlock(order=order, content=text)

    # ── 表格解析 ──────────────────────────────────────────────

    def _table_to_text(self, headers: List[str], rows: List[List[str]]) -> str:
        """表格 → markdown 格式文本。"""
        lines = [" | ".join(headers)]
        lines.append(" | ".join("---" for _ in headers))
        for row in rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    def _parse_table(self, table, start_order: int) -> Optional[TableBlock]:
        """解析 DOCX 表格 → TableBlock。"""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return None

        headers = rows_data[0]
        rows = [r for r in rows_data[1:] if any(c for c in r)]

        return TableBlock(
            headers=headers,
            rows=rows,
            content=self._table_to_text(headers, rows),
            order=start_order,
            metadata={
                "num_rows": len(headers) + len(rows),
                "num_cols": len(headers),
            },
        )

    # ── 图片提取 ──────────────────────────────────────────────

    def _extract_images(self, doc, doc_id: str, blocks: list, start_order: int):
        """从 DOCX 中提取图片 → ImageBlock（追加到 blocks）。"""
        order = start_order
        os.makedirs(self.image_output_dir, exist_ok=True)

        for para in doc.paragraphs:
            for run in para.runs:
                drawing_elements = run._element.findall(qn("w:drawing"))
                for drawing in drawing_elements:
                    blip = drawing.find(".//" + qn("a:blip"))
                    if blip is None:
                        continue
                    r_embed = blip.get(qn("r:embed"))
                    if r_embed is None:
                        continue

                    image_part = doc.part.related_parts.get(r_embed)
                    if image_part is None:
                        continue

                    img_path = os.path.join(
                        self.image_output_dir, f"{doc_id}_img{order}.png"
                    )
                    with open(img_path, "wb") as f:
                        f.write(image_part.blob)

                    blocks.append(
                        ImageBlock(
                            content="",
                            image_path=img_path,
                            alt_text=para.text[:100] or None,
                            order=order,
                        )
                    )
                    order += 1

    # ── 主解析流程 ────────────────────────────────────────────

    def _parse_impl(self, file_path: str) -> StructuredDocument:
        doc = Document(file_path)
        blocks: list = []
        order = 0
        para_idx = 0
        table_idx = 0

        for child in doc.element.body:
            tag = child.tag

            if tag == qn("w:p"):
                para = doc.paragraphs[para_idx]
                para_idx += 1
                block = self._parse_paragraph(para, order)
                if block:
                    blocks.append(block)
                    order += 1

            elif tag == qn("w:tbl") and self.extract_tables:
                tbl = doc.tables[table_idx]
                table_idx += 1
                block = self._parse_table(tbl, order)
                if block:
                    blocks.append(block)
                    order += 1

        doc_id = str(uuid.uuid4())

        if self.extract_images:
            self._extract_images(doc, doc_id, blocks, order)

        return StructuredDocument(
            doc_id=doc_id,
            source=file_path,
            blocks=blocks,
        )


# ═══════════════════════════════════════════════════════════════
#  法律文档 LegalDocxParser
# ═══════════════════════════════════════════════════════════════

class LegalDocxParser(DocxParser):
    """法律文档解析器，针对国家法律数据库的标准排版结构。

    法律文档特征：
    - 所有段落使用 Normal 样式，无原生 Heading
    - 标题为 22pt 加粗
    - 章节通过 `第X章`/`第X节` 模式标记
    - 条目通过 `第X条` 模式标记
    - 含目录 (TOC)，章节标题会重复出现
    """

    def __init__(
        self,
        extract_tables: bool = True,
        extract_images: bool = False,
        image_output_dir: str = "./docx_images",
    ):
        super().__init__(
            extract_tables=extract_tables,
            extract_images=extract_images,
            image_output_dir=image_output_dir,
        )

    # ── 自动检测 ──────────────────────────────────────────────

    def detect(self, file_path: str) -> bool:
        """检测是否为法律文档：扫描前 20 非空段落，章/节/条模式命中率 >30%。"""
        doc = Document(file_path)
        matches = 0
        non_empty = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            non_empty += 1
            if (
                _RE_CHAPTER.match(text)
                or _RE_SECTION.match(text)
                or _RE_ARTICLE.match(text)
            ):
                matches += 1
            if non_empty >= 20:
                break
        return non_empty > 0 and matches / non_empty > 0.3

    # ── 标题检测（覆盖） ──────────────────────────────────────

    def _get_heading_level(self, para) -> Optional[int]:
        """法律文档的标题检测：字号 + 正则模式。"""
        # 1. 22pt 文档标题 → level 1
        for run in para.runs:
            if run.font.size and run.font.size.pt >= 22:
                return 1

        text = para.text.strip()

        # 2. 序言 → level 2
        if _RE_PREFACE.match(text):
            return 2

        # 3. 第X章 → level 2
        if _RE_CHAPTER.match(text):
            return 2

        # 4. 附件X → level 2
        if _RE_APPENDIX.match(text):
            return 2

        # 5. 第X节 → level 3
        if _RE_SECTION.match(text):
            return 3

        return None

    # ── 段落解析（覆盖） ──────────────────────────────────────

    def _parse_paragraph(self, para, order: int):
        """法律文档段落解析：标题/条目/日期/普通。"""
        text = para.text.strip()
        if not text:
            return None

        # 标题
        heading_level = self._get_heading_level(para)
        if heading_level is not None:
            return HeadingBlock(order=order, content=text, level=heading_level)

        # 条目（第X条）
        article_match = _RE_ARTICLE.match(text)
        if article_match:
            article_num = chinese_to_arabic(article_match.group(1))
            return ParagraphBlock(
                order=order,
                content=text,
                metadata={"type": "article", "article_number": article_num},
            )

        # 颁布日期
        if text.startswith("（") and ("通过" in text or "公布" in text):
            return ParagraphBlock(
                order=order,
                content=text,
                metadata={"type": "enactment_date"},
            )

        # 普通段落（如条下的款、项）
        return ParagraphBlock(order=order, content=text)

    # ── 主解析流程（覆盖：TOC 去重 + 结构构建） ────────────────
    #
    # 法律文档的独特挑战：
    #   1. 所有段落都是 Normal 样式，没有原生 Heading，不能依赖样式判断标题
    #   2. 文档包含目录(TOC)，章节标题在目录和正文中重复出现，必须跳过目录
    #   3. 需要通过正则模式（第X章/第X节/第X条）+ 字号来识别结构
    #
    # 解析策略：
    #   遍历 docx 的 XML body，同步维护 para_idx/table_idx 两个指针，
    #   因为 python-docx 的 .paragraphs/.tables 是独立列表，
    #   而 body 中的 w:p（段落）和 w:tbl（表格）是交错排列的，
    #   必须用索引来追踪当前处理到哪个段落/哪个表格。
    #
    #   与此同时，维护一个 structure 字典记录文档的逻辑结构
    #   （章→节→条），供后续检索利用层级信息。

    def _parse_impl(self, file_path: str) -> StructuredDocument:
        doc = Document(file_path)
        blocks: list = []
        order = 0

        # ── 双指针：同步遍历 body 中的 w:p 和 w:tbl ──
        # python-docx 的 doc.paragraphs 和 doc.tables 是顺序列表，
        # 而 doc.element.body 中的子标签是交错排列的（p / tbl 交替出现）。
        # para_idx 和 table_idx 分别跟踪当前处理到第几个段落/表格。
        para_idx = 0
        table_idx = 0

        # ── TOC 状态管理 ──
        # 法律文档通常以目录开头，目录中包含了所有章节标题。
        # 如果不过滤，这些标题会和正文中的标题重复出现，
        # 导致 chunker 产生重复的 heading_path 或空的 chunk。
        #
        # 检测逻辑：
        #   - 遇到 "目录" 两个字 → 进入 TOC mode (in_toc = True)
        #   - TOC mode 下所有内容都跳过（不生成 block）
        #   - 遇到第一个空行 → 退出 TOC mode
        #   - 注意：TOC 自身的 "目录" 二字保留（标记 type=toc 存入 blocks）
        in_toc = False

        # ── 结构目录（用于 metadata） ──
        # 记录文档的章→节→条层级，不直接影响 block 生成，
        # 但会存入 StructuredDocument.metadata 供后续检索/展示使用。
        #
        # 结构示例：
        # {
        #   "chapters": [
        #     {"num": 1, "title": "总则", "sections": [], "articles": [1,2,3]},
        #     {"num": 2, "title": "分则", "sections": [
        #       {"num": 1, "title": "第一节", "articles": [4,5,6]},
        #     ], "articles": []},
        #   ],
        #   "preamble_articles": [1]  # 第一章之前的条（如"第一条 为了..."）
        # }
        structure: dict = {"chapters": [], "preamble_articles": []}
        current_chapter: Optional[dict] = None   # 当前所在的章，用于挂载节和条
        current_section: Optional[dict] = None    # 当前所在的节，用于挂载条

        # ── 遍历 XML body，段落和表格交替处理 ──
        # 为什么不用 doc.paragraphs 直接遍历？
        # 因为表格中的段落也会出现在 doc.paragraphs 中，无法区分。
        # 遍历 body 可以准确区分当前是段落(w:p)还是表格(w:tbl)。
        for child in doc.element.body:
            tag = child.tag

            # ── 段落处理 ──────────────────────────────────
            if tag == qn("w:p"):
                # 用 para_idx 从 doc.paragraphs 中取对应的段落对象
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                # _parse_paragraph 返回 None 表示空段落或纯空白
                block = self._parse_paragraph(para, order)

                if block is None:
                    # 空行/空白段落 — 在 TOC mode 中充当「结束标记」
                    # （TOC 之后通常跟一个空行，正好用来退出 TOC mode）
                    if in_toc:
                        in_toc = False
                    continue

                # ── TOC 检测入口 ──
                # 一旦匹配"目录"二字，进入 TOC mode，
                # 后续所有内容跳过，直到遇到空行。
                # 注意：toc_headings 集合在这里声明但未使用，
                # 留作将来可能需要记录 TOC 中的标题去重使用。
                if _RE_TOC.match(text):
                    in_toc = True
                    block.metadata["type"] = "toc"
                    blocks.append(block)
                    order += 1
                    continue

                # ── TOC mode：跳过目录内容 ──
                # 目录中会包含"第X章"等标题，与正文重复，
                # 直接丢弃（不生成 block），避免下游出现空 chunk。
                if in_toc:
                    continue

                # ── 结构目录构建（不影响 block 本身） ──
                # 这里只做一件事：追踪当前解析到哪一章/哪一节，
                # 把遇到的条号挂到对应的章/节下，形成结构索引。
                # 这个 structure 最终存入 metadata 供检索用。
                ch_match = _RE_CHAPTER.match(text)
                sec_match = _RE_SECTION.match(text)
                art_match = _RE_ARTICLE.match(text)

                if ch_match:
                    # 遇到新章：重置 current_chapter，清空 current_section
                    ch_num = chinese_to_arabic(ch_match.group(1))
                    current_chapter = {
                        "num": ch_num,
                        "title": ch_match.group(2).strip(),
                        "sections": [],
                        "articles": [],
                    }
                    current_section = None
                    structure["chapters"].append(current_chapter)

                elif sec_match:
                    # 遇到新节：挂到 current_chapter.sections 下
                    # 注意：如果节出现在所有章之前（罕见），
                    # current_chapter 为 None，此时节信息会丢失。
                    sec_num = chinese_to_arabic(sec_match.group(1))
                    current_section = {
                        "num": sec_num,
                        "title": sec_match.group(2).strip(),
                        "articles": [],
                    }
                    if current_chapter:
                        current_chapter["sections"].append(current_section)

                elif art_match:
                    # 遇到新条：挂到当前节 → 当前章 → preamble_articles
                    # 优先级：节 > 章 > 前言
                    art_num = chinese_to_arabic(art_match.group(1))
                    if current_section:
                        current_section["articles"].append(art_num)
                    elif current_chapter:
                        current_chapter["articles"].append(art_num)
                    else:
                        structure["preamble_articles"].append(art_num)

                # 无论是否匹配章/节/条，所有非空段落都生成 block
                # （结构构建仅用于 metadata 索引，不影响 block 生成）
                blocks.append(block)
                order += 1

            # ── 表格处理 ──────────────────────────────────
            elif tag == qn("w:tbl") and self.extract_tables:
                tbl = doc.tables[table_idx]
                table_idx += 1
                block = self._parse_table(tbl, order)
                if block:
                    blocks.append(block)
                    order += 1

        doc_id = str(uuid.uuid4())

        if self.extract_images:
            self._extract_images(doc, doc_id, blocks, order)

        return StructuredDocument(
            doc_id=doc_id,
            source=file_path,
            blocks=blocks,
            metadata={
                "doc_type": "legal",
                "structure": structure,
            },
        )
